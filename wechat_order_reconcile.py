"""
离线版：微信单号群记录追加 + 订单结果对比
-------------------------------------------------
功能：
1) 读取微信导出xlsx（B列时间，E列内容），提取有效单号
2) 追加/合并到微信汇总库 wechat_shipment_log.xlsx
3) 将订单结果与微信汇总库对比，输出发货核对结果

运行示例（推荐，最简）：
python wechat_order_reconcile.py
# 然后按提示回车确认默认路径

全自动模式：
python wechat_order_reconcile.py --auto
# 自动选 data 目录中最近的微信导出与订单结果文件

手动指定模式：
python wechat_order_reconcile.py --wechat "data/群聊_单号群.xlsx" --orders "data/orders_result_20260308_100629(1).xlsx"
"""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


def init_console_encoding():
    """避免 Windows GBK 控制台打印 emoji 时抛出编码异常。"""
    for stream_name in ("stdout", "stderr"):
        stream = getattr(sys, stream_name, None)
        if stream and hasattr(stream, "reconfigure"):
            try:
                stream.reconfigure(encoding="utf-8", errors="replace")
            except Exception:
                pass


ORDER_PATTERN = re.compile(r"\d{4,}(?:-\d+)*")
DOCX_ORDER_PATTERN = re.compile(r"(?<!\d)[1-9]\d{4,7}(?:-\d{1,3}){0,3}(?!\d)")
DOCX_TIME_PATTERN = re.compile(
    r"(?:(昨天)\s*)?(\d{1,2})月(\d{1,2})日(?:星期[一二三四五六日])?\s*(\d{1,2}):(\d{2})"
)
DOCX_SIMPLE_TIME_PATTERN = re.compile(r"(昨天)\s*(\d{1,2}):(\d{2})")


def extract_base_order(order_no: str) -> str:
    m = re.match(r"^(\d+)", str(order_no or "").strip())
    return m.group(1) if m else ""


def normalize_text(val) -> str:
    if val is None:
        return ""
    s = str(val).strip()
    if s.lower() in {"nan", "none"}:
        return ""
    return s


def extract_orders_from_content(content: str) -> list[str]:
    """
    从一条群消息中提取多个单号。
    例如："99759 99721一起发" -> ["99759", "99721"]
    """
    s = normalize_text(content)
    if not s:
        return []
    if s.startswith("[") and s.endswith("]"):
        return []

    matches = ORDER_PATTERN.findall(s)
    # 去重但保持顺序
    seen = set()
    result = []
    for m in matches:
        if m not in seen:
            seen.add(m)
            result.append(m)
    return result


def is_docx_file(path: Path) -> bool:
    return path.suffix.lower() == ".docx"


def iter_docx_lines(doc_path: Path):
    if not zipfile.is_zipfile(doc_path):
        raise ValueError(
            f"文件不是有效的 DOCX（zip 容器）: {doc_path}\n"
            "请确认文件后缀为 .docx，且不是 .doc 或已损坏文件。"
        )
    try:
        doc = Document(str(doc_path))
    except PackageNotFoundError as exc:
        raise FileNotFoundError(f"无法打开 DOCX 文件: {doc_path}") from exc

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            yield text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if text:
                    yield text


def extract_orders_from_docx(doc_path: Path) -> pd.DataFrame:
    lines = list(iter_docx_lines(doc_path))
    rows = []
    last_time_text = ""

    for raw in lines:
        text = normalize_text(raw)
        if not text:
            continue

        last_time_text = parse_docx_time_text(text, last_time_text)

        matches = DOCX_ORDER_PATTERN.findall(text)
        if not matches:
            continue

        for order_no in matches:
            rows.append(
                {
                    "chat_time": last_time_text,
                    "content": text,
                    "order_no": order_no,
                }
            )

    if not rows:
        return pd.DataFrame(columns=["chat_time", "content", "order_no", "base_order_no", "source_file", "import_time"])

    df = pd.DataFrame(rows)
    df["base_order_no"] = df["order_no"].map(extract_base_order)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    df["source_file"] = str(doc_path.name)
    df["import_time"] = now
    df["chat_time"] = df["chat_time"].map(normalize_text)
    df["content"] = df["content"].map(normalize_text)
    return df[["chat_time", "content", "order_no", "base_order_no", "source_file", "import_time"]]


def parse_docx_time_text(text: str, last_time: str) -> str:
    normalized = text.replace("：", ":")
    m = DOCX_TIME_PATTERN.search(normalized)
    if m:
        is_yesterday, month, day, hour, minute = m.groups()
        dt = build_docx_datetime(is_yesterday, month, day, hour, minute)
        return dt.strftime("%Y-%m-%d %H:%M:%S") if dt else last_time

    m2 = DOCX_SIMPLE_TIME_PATTERN.search(normalized)
    if m2:
        is_yesterday, hour, minute = m2.groups()
        dt = build_docx_datetime(is_yesterday, None, None, hour, minute)
        return dt.strftime("%Y-%m-%d %H:%M:%S") if dt else last_time

    return last_time


def build_docx_datetime(is_yesterday: str | None, month: str | None, day: str | None, hour: str, minute: str):
    now = datetime.now()
    target_date = now.date()

    if month and day:
        try:
            target_date = datetime(now.year, int(month), int(day)).date()
        except ValueError:
            target_date = now.date()
    if is_yesterday:
        target_date = (datetime.combine(target_date, datetime.min.time()) - timedelta(days=1)).date()

    try:
        dt = datetime.combine(target_date, datetime.min.time()).replace(hour=int(hour), minute=int(minute))
    except ValueError:
        return None
    return dt


def load_wechat_orders(wechat_xlsx: Path) -> pd.DataFrame:
    # 按你给的固定规则：B列=时间，E列=单号内容
    raw = pd.read_excel(wechat_xlsx, header=None, dtype=str)

    if raw.shape[1] < 5:
        raise RuntimeError("微信导出表列数不足，无法读取E列单号内容")

    df = pd.DataFrame(
        {
            "chat_time": raw.iloc[:, 1].map(normalize_text),
            "content": raw.iloc[:, 4].map(normalize_text),
        }
    )

    # 一条消息可能含多个单号：先提取列表，再 explode 成多行
    df["order_list"] = df["content"].map(extract_orders_from_content)
    df = df.explode("order_list", ignore_index=True)
    df["order_no"] = df["order_list"].map(normalize_text)
    df = df[df["order_no"] != ""].copy()

    if df.empty:
        return pd.DataFrame(
            columns=[
                "chat_time",
                "content",
                "order_no",
                "base_order_no",
                "source_file",
                "import_time",
            ]
        )

    df["base_order_no"] = df["order_no"].map(extract_base_order)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    df["source_file"] = str(wechat_xlsx.name)
    df["import_time"] = now

    return df[["chat_time", "content", "order_no", "base_order_no", "source_file", "import_time"]]


def load_input_orders(input_path: Path, source_type: str) -> pd.DataFrame:
    if source_type == "docx":
        return extract_orders_from_docx(input_path)
    return load_wechat_orders(input_path)


def build_clean_orders(raw_imports: pd.DataFrame) -> pd.DataFrame:
    if raw_imports.empty:
        return pd.DataFrame(
            columns=["order_no", "base_order_no", "first_seen", "last_seen", "hit_count"]
        )

    tmp = raw_imports.copy()
    tmp["chat_time_dt"] = pd.to_datetime(tmp["chat_time"], errors="coerce")
    tmp["fallback_dt"] = pd.to_datetime(tmp["import_time"], errors="coerce")
    tmp["time_for_stats"] = tmp["chat_time_dt"].fillna(tmp["fallback_dt"])

    # 你要求按“单号去重”：order_no 唯一
    grouped = (
        tmp.groupby(["order_no", "base_order_no"], as_index=False)
        .agg(
            first_seen=("time_for_stats", "min"),
            last_seen=("time_for_stats", "max"),
            hit_count=("order_no", "count"),
        )
        .sort_values(["order_no"])
    )

    grouped["first_seen"] = grouped["first_seen"].dt.strftime("%Y-%m-%d %H:%M:%S").fillna("")
    grouped["last_seen"] = grouped["last_seen"].dt.strftime("%Y-%m-%d %H:%M:%S").fillna("")
    return grouped


def load_sheet_if_exists(path: Path, sheet_name: str) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    try:
        return pd.read_excel(path, sheet_name=sheet_name, dtype=str)
    except Exception:
        return pd.DataFrame()


def append_to_wechat_log(wechat_df_new: pd.DataFrame, log_path: Path) -> tuple[pd.DataFrame, pd.DataFrame]:
    old_raw = load_sheet_if_exists(log_path, "raw_imports")

    if old_raw.empty:
        merged_raw = wechat_df_new.copy()
    else:
        merged_raw = pd.concat([old_raw, wechat_df_new], ignore_index=True)

    for col in ["chat_time", "content", "order_no", "base_order_no", "source_file", "import_time"]:
        if col not in merged_raw.columns:
            merged_raw[col] = ""

    merged_raw = merged_raw[["chat_time", "content", "order_no", "base_order_no", "source_file", "import_time"]]

    clean_orders = build_clean_orders(merged_raw)

    log_path.parent.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(log_path, engine="openpyxl") as writer:
        merged_raw.to_excel(writer, sheet_name="raw_imports", index=False)
        clean_orders.to_excel(writer, sheet_name="clean_orders", index=False)

    return merged_raw, clean_orders


def detect_order_column(df: pd.DataFrame) -> Optional[str]:
    priority = [
        "订单编号(oId)",
        "oId",
        "订单编号",
        "查询编号",
        "原始输入",
    ]
    cols = [str(c).strip() for c in df.columns]

    for p in priority:
        if p in cols:
            return p

    for c in cols:
        if "编号" in c or c.lower() == "oid":
            return c

    return None


def detect_status_column(df: pd.DataFrame) -> Optional[str]:
    priority = ["订单状态", "状态文字", "状态"]
    cols = [str(c).strip() for c in df.columns]

    for p in priority:
        if p in cols:
            return p

    for c in cols:
        if "状态" in c:
            return c

    return None


def status_is_done(status_text: str) -> bool:
    s = normalize_text(status_text)
    return ("打包" in s and "完成" in s) or ("已完成" in s)


def compare_orders(orders_xlsx: Path, clean_orders: pd.DataFrame) -> pd.DataFrame:
    orders = pd.read_excel(orders_xlsx, dtype=str)
    orders.columns = [str(c).strip() for c in orders.columns]

    col_order = detect_order_column(orders)
    if not col_order:
        raise RuntimeError(f"订单文件找不到编号列，当前列: {list(orders.columns)}")

    col_status = detect_status_column(orders)

    clean = clean_orders.copy()
    for c in ["order_no", "base_order_no", "first_seen", "last_seen", "hit_count"]:
        if c not in clean.columns:
            clean[c] = ""

    exact_map = {normalize_text(r["order_no"]): r for _, r in clean.iterrows()}
    base_group = clean.groupby("base_order_no", as_index=False).agg(
        first_seen=("first_seen", "min"),
        last_seen=("last_seen", "max"),
        hit_count=("hit_count", "sum"),
    )
    base_map = {normalize_text(r["base_order_no"]): r for _, r in base_group.iterrows()}

    out_rows = []

    for _, row in orders.iterrows():
        order_no = normalize_text(row.get(col_order, ""))
        base_no = extract_base_order(order_no)
        status_text = normalize_text(row.get(col_status, "")) if col_status else ""

        exact_hit = exact_map.get(order_no)
        base_hit = base_map.get(base_no) if base_no else None

        if exact_hit is not None:
            wechat_match = "精确命中"
            first_seen = normalize_text(exact_hit.get("first_seen", ""))
            last_seen = normalize_text(exact_hit.get("last_seen", ""))
            hit_count = normalize_text(exact_hit.get("hit_count", ""))
        elif base_hit is not None:
            wechat_match = "基础号命中"
            first_seen = normalize_text(base_hit.get("first_seen", ""))
            last_seen = normalize_text(base_hit.get("last_seen", ""))
            hit_count = normalize_text(base_hit.get("hit_count", ""))
        else:
            wechat_match = "未命中"
            first_seen = ""
            last_seen = ""
            hit_count = ""

        done = status_is_done(status_text)

        if done and wechat_match != "未命中":
            final_judgement = "已发货"
        elif done and wechat_match == "未命中":
            final_judgement = "待核实(已完成未见群记录)"
        elif (not done) and wechat_match != "未命中":
            final_judgement = "异常(未完成但群有记录)"
        else:
            final_judgement = "正常未发货"

        out = dict(row)
        out.update(
            {
                "匹配订单号": order_no,
                "基础号": base_no,
                "微信匹配": wechat_match,
                "微信首现时间": first_seen,
                "微信末现时间": last_seen,
                "微信出现次数": hit_count,
                "最终判定": final_judgement,
            }
        )
        out_rows.append(out)

    return pd.DataFrame(out_rows)


def save_reconcile_result(result_df: pd.DataFrame, clean_orders: pd.DataFrame, output_path: Path):
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        result_df.to_excel(writer, sheet_name="核对结果", index=False)
        clean_orders.to_excel(writer, sheet_name="微信单号汇总", index=False)


def find_latest_xlsx_by_keywords(data_dir: Path, keywords: list[str]) -> Optional[Path]:
    if not data_dir.exists():
        return None

    files = [p for p in data_dir.glob("*.xlsx") if p.is_file()]
    if not files:
        return None

    def score(p: Path):
        name = p.name.lower()
        hit = sum(1 for k in keywords if k in name)
        return (hit, p.stat().st_mtime)

    files.sort(key=score, reverse=True)
    best = files[0]
    if score(best)[0] <= 0:
        return None
    return best


def find_latest_docx_by_keywords(data_dir: Path, keywords: list[str]) -> Optional[Path]:
    if not data_dir.exists():
        return None

    files = [p for p in data_dir.glob("*.docx") if p.is_file()]
    if not files:
        return None

    def score(p: Path):
        name = p.name.lower()
        hit = sum(1 for k in keywords if k in name)
        return (hit, p.stat().st_mtime)

    files.sort(key=score, reverse=True)
    best = files[0]
    if score(best)[0] <= 0:
        return None
    return best


def prompt_path(label: str, default_path: Path) -> Path:
    text = input(f"{label}（回车使用默认）\n> {default_path}\n> ").strip()
    if not text:
        return default_path
    text = text.strip().strip('"').strip("'")
    return Path(text)


def parse_args():
    parser = argparse.ArgumentParser(description="微信单号追加与订单发货核对（离线）")
    parser.add_argument("--wechat", default="", help="微信导出xlsx路径（可省略，交互选择）")
    parser.add_argument("--orders", default="", help="订单结果xlsx路径（可省略，交互选择）")
    parser.add_argument("--log", default="data/wechat_shipment_log.xlsx", help="微信汇总库路径")
    parser.add_argument(
        "--output",
        default=f"data/reconcile_result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
        help="核对结果输出路径",
    )
    parser.add_argument(
        "--auto",
        action="store_true",
        help="自动使用data目录最近文件（微信+订单），不进入交互",
    )
    parser.add_argument(
        "--no-interactive",
        action="store_true",
        help="禁用交互输入，路径缺失时直接报错",
    )
    parser.add_argument(
        "--source",
        choices=["auto", "wechat", "docx"],
        default="auto",
        help="输入来源类型（自动识别/微信导出/docx）",
    )
    return parser.parse_args()


def detect_input_source(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return "docx"
    if path.suffix.lower() in {".xlsx", ".xlsm"}:
        return "wechat"
    return "wechat"


def main():
    init_console_encoding()
    args = parse_args()

    log_path = Path(args.log)
    output_path = Path(args.output)

    data_dir = Path("data")
    auto_wechat = find_latest_xlsx_by_keywords(data_dir, ["群聊", "单号", "wechat"])
    auto_docx = find_latest_docx_by_keywords(data_dir, ["docx", "微信", "聊天", "导出"])
    auto_orders = find_latest_xlsx_by_keywords(data_dir, ["orders_result", "orders", "result", "核对", "订单"])

    input_path = Path(args.wechat) if args.wechat else (auto_wechat or auto_docx)
    orders_path = Path(args.orders) if args.orders else auto_orders

    should_prompt = (not args.auto) and (not args.no_interactive) and (
        not args.wechat or not args.orders
    )

    if should_prompt:
        print("\n[INFO] 微信单号核对（简化输入）")
        if input_path is None:
            input_path = Path("data/群聊_单号群.xlsx")
        if orders_path is None:
            orders_path = Path("data/orders_result.xlsx")

        input_path = prompt_path("微信导出xlsx路径/Docx路径", input_path)
        orders_path = prompt_path("订单结果xlsx路径", orders_path)

    if input_path is None or not Path(input_path).exists():
        raise FileNotFoundError(f"微信/Docx 文件不存在: {input_path}")
    if orders_path is None or not Path(orders_path).exists():
        raise FileNotFoundError(f"订单文件不存在: {orders_path}")

    input_path = Path(input_path)
    orders_path = Path(orders_path)

    if args.source == "auto":
        source_type = detect_input_source(input_path)
    else:
        source_type = args.source

    source_label = "DOCX" if source_type == "docx" else "微信导出"

    print(f"[STEP 1/4] 读取{source_label}: {input_path}")
    wechat_new = load_input_orders(input_path, source_type)
    print(f"[INFO] 提取到有效单号 {len(wechat_new)} 条")

    print(f"[STEP 2/4] 追加到微信汇总库: {log_path}")
    _, clean_orders = append_to_wechat_log(wechat_new, log_path)
    print(f"[INFO] 汇总后唯一单号 {len(clean_orders)} 条")

    print(f"[STEP 3/4] 对比订单文件: {orders_path}")
    reconcile_df = compare_orders(orders_path, clean_orders)
    print(f"[INFO] 生成核对记录 {len(reconcile_df)} 条")

    print(f"[STEP 4/4] 输出结果: {output_path}")
    save_reconcile_result(reconcile_df, clean_orders, output_path)

    print("\n[OK] 完成：已输出核对结果并更新微信汇总库")
    print("\n[INFO] 提示：下次可直接运行 `python wechat_order_reconcile.py --auto` 自动选data目录最新文件。")


if __name__ == "__main__":
    main()
