"""
棉花果(58mhg.com) 订单进度批量查询爬虫
======================================
使用方法:
  1. 在 CONFIG 区域更新 Cookie（过期后重新复制替换）
  2. 准备好 Excel 文件，填入用户名和订单编号
  3. 运行: python mhg_batch_query.py

依赖安装: pip install requests openpyxl pandas
"""

import os
import sys
import argparse
import requests
import pandas as pd
import openpyxl
import re
import hashlib
import getpass
import json
import zipfile
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from datetime import datetime, timedelta
import time
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


def init_console_encoding():
    """避免 Windows GBK 控制台打印 emoji 时抛出编码异常。"""
    for stream_name in ("stdout", "stderr"):
        stream = getattr(sys, stream_name, None)
        if stream and hasattr(stream, "reconfigure"):
            try:
                stream.reconfigure(encoding="utf-8", errors="replace")
            except Exception:
                pass


# ============================================================
#  ★ 配置区 —— 只需修改这里 ★
# ============================================================

def load_env_file(env_path: str = ".env"):
    if not os.path.isfile(env_path):
        return
    try:
        with open(env_path, "r", encoding="utf-8") as f:
            for raw_line in f:
                line = raw_line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                key, value = line.split("=", 1)
                key = key.strip()
                value = value.strip().strip('"').strip("'")
                if key and key not in os.environ:
                    os.environ[key] = value
    except Exception:
        pass


def get_env(name: str, default: str = "") -> str:
    return str(os.getenv(name, default) or "").strip()


def get_env_float(name: str, default: float) -> float:
    raw = get_env(name, "")
    if not raw:
        return default
    try:
        return float(raw)
    except ValueError:
        return default


load_env_file()

CONFIG = {
    # Cookie（从浏览器F12复制，过期后重新复制替换即可）
    "cookie": get_env("MHG_COOKIE"),

    # 输入Excel文件路径
    "input_excel": get_env("MHG_INPUT_EXCEL", ""),

    # 输出Excel文件路径
    "output_excel": get_env(
        "MHG_OUTPUT_EXCEL",
        f"orders_result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
    ),

    # 每次请求间隔秒数（避免请求过快）
    "request_delay": get_env_float("MHG_REQUEST_DELAY", 1.0),

    # 查询日期范围（建议放宽，确保能查到所有订单）
    "start_date": get_env("MHG_START_DATE", "2024-01-01"),
    "end_date": get_env("MHG_END_DATE", "2026-12-31"),

    # 微信导出及汇总库配置
    "wechat_mode": get_env("MHG_WECHAT_MODE", "append"),  # append / merge_new
    "wechat_export_excel": get_env("MHG_WECHAT_EXPORT_EXCEL", ""),  # 微信导出xlsx路径（留空表示本次跳过微信比对）
    "wechat_repo_excel": get_env("MHG_WECHAT_REPO_EXCEL", "wechat_shipment_log.xlsx"),  # 微信汇总库
}
# ============================================================


API_GET_LIST  = "https://58mhg.com/index/orders/getList"
API_GET_COUNT = "https://58mhg.com/index/orders/getOrdersCount"
API_STEP_PROGRESS = "https://58mhg.com/index/orders/stepProgress"

AUTH_CONFIG_FILE = "auth_config.json"
AUTH_PEPPER = "MHG@2026#scan-query#pepper"

# Excel 行颜色
COLOR_NOT_FOUND = "FCE4D6"  # 红：查无此单
COLOR_PENDING   = "DDEBF7"  # 蓝：待开始（process=0，订单存在但未生产）
COLOR_IN_PROG   = "FFF2CC"  # 黄：生产进行中
COLOR_DONE      = "E2EFDA"  # 绿：已完成


def ask(prompt: str, default: str = "") -> str:
    val = input(prompt).strip()
    return val if val else default


def clean_cli_path(text: str) -> str:
    """清理命令行输入路径（兼容拖拽文件带引号）。"""
    val = str(text or "").strip()
    if (val.startswith('"') and val.endswith('"')) or (val.startswith("'") and val.endswith("'")):
        val = val[1:-1].strip()
    return val


def prompt_runtime_paths():
    """
    运行时输入路径：
    - 支持直接回车使用默认值
    - 支持 Windows 拖拽文件到终端（自动去除两侧引号）
    """
    default_in = CONFIG["input_excel"]
    default_out = CONFIG["output_excel"]
    default_wechat = CONFIG.get("wechat_export_excel", "")
    default_repo = CONFIG.get("wechat_repo_excel", "wechat_shipment_log.xlsx")
    default_mode = CONFIG.get("wechat_mode", "append")

    print("\n请确认本次文件路径（支持把 Excel 直接拖到窗口）：")
    in_raw = input(f"输入Excel路径（回车用默认）\n> {default_in}\n> ").strip()
    out_raw = input(f"输出Excel路径（回车用默认）\n> {default_out}\n> ").strip()

    print("\n微信导出表（可选，不填则跳过发货比对）：")
    wechat_raw = input(f"微信导出xlsx路径（回车可跳过）\n> {default_wechat}\n> ").strip()
    repo_raw = input(f"微信汇总库路径（回车用默认）\n> {default_repo}\n> ").strip()
    mode_raw = input(f"微信导入模式 append/merge_new（回车用默认）\n> {default_mode}\n> ").strip().lower()

    input_excel = clean_cli_path(in_raw) if in_raw else default_in
    output_excel = clean_cli_path(out_raw) if out_raw else default_out
    wechat_excel = clean_cli_path(wechat_raw) if wechat_raw else default_wechat
    wechat_repo = clean_cli_path(repo_raw) if repo_raw else default_repo
    wechat_mode = mode_raw if mode_raw in ("append", "merge_new") else default_mode

    return input_excel, output_excel, wechat_excel, wechat_repo, wechat_mode


def sha256_hex(text: str) -> str:
    return hashlib.sha256(str(text or "").encode("utf-8")).hexdigest()


def hash_password(password: str) -> str:
    mixed = f"{AUTH_PEPPER}{password}{AUTH_PEPPER}"
    return sha256_hex(mixed)


def load_auth_hash(config_path: str) -> str:
    if not os.path.isfile(config_path):
        return ""
    try:
        with open(config_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return str(data.get("password_sha256", "")).strip().lower()
    except Exception:
        return ""


def check_password_or_exit(max_try: int = 3):
    """启动密码校验：读取同目录 auth_config.json 的 password_sha256。"""
    current_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
    config_path = os.path.join(current_dir, AUTH_CONFIG_FILE)
    target_hash = load_auth_hash(config_path)

    if not target_hash:
        print(f"\n[WARN] 未读取到密码配置: {config_path}")
        print("  请先运行 set_password.py 设置密码后再使用。")
        raise SystemExit(1)

    print("\n[AUTH] 请输入启动密码")
    for i in range(1, max_try + 1):
        try:
            pwd = getpass.getpass("密码: ")
        except Exception:
            pwd = input("密码: ").strip()

        if hash_password(pwd).lower() == target_hash:
            print("[OK] 密码正确，开始运行。\n")
            return

        left = max_try - i
        if left > 0:
            print(f"[ERR] 密码错误，还可尝试 {left} 次")
            time.sleep(1.2)

    print("❌ 密码连续错误，程序已退出。")
    raise SystemExit(1)


def build_headers():
    return {
        "Content-Type": "application/json",
        "Cookie": CONFIG["cookie"],
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Referer": "https://58mhg.com/console/",
        "Origin": "https://58mhg.com",
    }


def build_payload(keyword: str, page_index: int = 1, page_size: int = 20):
    return {
        "pageIndex": page_index,
        "pageSize": page_size,
        "pageItem": 1,
        "keyWord": str(keyword),
        "showMaterial": True,
        "withDepartment": False,
        "startDate": CONFIG["start_date"],
        "endDate": CONFIG["end_date"],
        "startOut": "",
        "endOut": "",
        "key": "like",
        "isBuild": True,
    }


def process_to_status(process: float):
    """process 数值 → (状态文字, 行颜色)；作为兜底判定。"""
    if process >= 0.75:
        return "已完成", COLOR_DONE
    elif process == 0.0:
        return "待开始", COLOR_PENDING
    elif process == 0.25:
        return "生产中 25%", COLOR_IN_PROG
    elif process == 0.5:
        return "生产中 50%", COLOR_IN_PROG
    else:
        return f"生产中 {int(process*100)}%", COLOR_IN_PROG


def parse_progress_value(raw):
    """将 progress 安全转为 float，失败返回 0。"""
    try:
        return float(raw)
    except Exception:
        return 0.0


def is_packaging_done(step_progress_data: list) -> bool:
    """
    stepProgress 规则：
    - 只要工序 name 包含“打包”
    - 且该工序 progress == 1
    即判定订单已完成。
    """
    if not isinstance(step_progress_data, list):
        return False

    for step in step_progress_data:
        if not isinstance(step, dict):
            continue
        name = str(step.get("name", "")).strip()
        if "打包" not in name:
            continue
        progress = parse_progress_value(step.get("progress"))
        if progress >= 1:
            return True

    return False


def get_step_progress(session: requests.Session, headers: dict, order_internal_id: str):
    """按订单内部ID拉取 stepProgress 数据。"""
    if not order_internal_id:
        return []

    try:
        resp = session.post(
            API_STEP_PROGRESS,
            json={"id": order_internal_id},
            headers=headers,
            timeout=15,
        )
        data = resp.json()
        if data.get("code") != 200:
            return []
        steps = data.get("data", [])
        return steps if isinstance(steps, list) else []
    except Exception:
        return []


def decide_status_by_step_or_process(session: requests.Session, headers: dict, order: dict):
    """
    状态判定优先级：
    1) stepProgress 中“打包”工序 progress=1 -> 已完成
    2) 否则按 process 兜底
    返回: (status_text, row_color, process)
    """
    order_internal_id = str(order.get("id", "")).strip()
    steps = get_step_progress(session, headers, order_internal_id)

    raw_process = order.get("process")
    process = parse_progress_value(raw_process)

    if is_packaging_done(steps):
        return "已完成(打包完成)", COLOR_DONE, process

    status_text, row_color = process_to_status(process)
    return status_text, row_color, process


def normalize_query_id(raw_text: str) -> str:
    """
    规范化查询编号：
    1) 若输入是路径（如 网盘/777/666/777），只取第二层（777）
    2) 若编号包含分表后缀（如 99556-1），查询时只取基础编号（99556）
    3) 兼容 / 和 \\ 分隔符
    """
    text = str(raw_text or "").strip()
    if not text:
        return ""

    # 统一路径分隔符，并去掉空段
    parts = [p.strip() for p in text.replace("\\", "/").split("/") if p.strip()]

    # 按你的规则：只取第二层（例如 网盘/777/666/777 -> 777）
    candidate = parts[1] if len(parts) >= 2 else text

    # 分表编号归一：99556-1 -> 99556
    m = re.match(r"^(\d+)", candidate)
    if m:
        return m.group(1)

    # 非数字前缀则按原值
    return candidate


def normalize_order_code(raw_text: str) -> str:
    """标准化订单号（保留数字和-，统一小写）。"""
    text = str(raw_text or "").strip().lower()
    if not text:
        return ""
    text = text.replace("—", "-")
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"[^0-9\-]", "", text)
    return text.strip("-")


def extract_order_codes_from_text(raw_text: str) -> list:
    """从微信内容中提取订单号片段，如 99641 / 99641-2 / 99441-1-2。"""
    text = str(raw_text or "")
    if not text:
        return []
    matches = re.findall(r"\d+(?:-\d+)*", text)
    out = []
    seen = set()
    for m in matches:
        code = normalize_order_code(m)
        if not code or code in seen:
            continue
        seen.add(code)
        out.append(code)
    return out


def get_base_code(code: str) -> str:
    """基础号：99641-2 -> 99641。"""
    m = re.match(r"^(\d+)", str(code or ""))
    return m.group(1) if m else ""


def is_wechat_noise(text: str) -> bool:
    """判断微信导出内容是否为噪音消息。"""
    t = str(text or "").strip()
    if not t:
        return True
    noise_tokens = {"[图片]", "[其他消息]", "[动画表情]", "[语音]", "[视频]", "[链接]"}
    return t in noise_tokens


def iter_docx_lines(doc_path: str):
    if not zipfile.is_zipfile(doc_path):
        raise ValueError(
            f"文件不是有效的 DOCX（zip 容器）: {doc_path}\n"
            "请确认文件后缀为 .docx，且不是 .doc 或已损坏文件。"
        )
    try:
        doc = Document(str(doc_path))
    except PackageNotFoundError as exc:
        raise FileNotFoundError(f"无法打开 DOCX 文件: {doc_path}") from exc

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            yield text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if text:
                    yield text


def parse_docx_time_text(text: str, last_time: str) -> str:
    normalized = text.replace("：", ":")
    m = re.search(r"(?:(昨天)\s*)?(\d{1,2})月(\d{1,2})日(?:星期[一二三四五六日])?\s*(\d{1,2}):(\d{2})", normalized)
    if m:
        is_yesterday, month, day, hour, minute = m.groups()
        dt = build_docx_datetime(is_yesterday, month, day, hour, minute)
        return dt.strftime("%Y-%m-%d %H:%M:%S") if dt else last_time

    m2 = re.search(r"(昨天)\s*(\d{1,2}):(\d{2})", normalized)
    if m2:
        is_yesterday, hour, minute = m2.groups()
        dt = build_docx_datetime(is_yesterday, None, None, hour, minute)
        return dt.strftime("%Y-%m-%d %H:%M:%S") if dt else last_time

    return last_time


def build_docx_datetime(is_yesterday: str | None, month: str | None, day: str | None, hour: str, minute: str):
    now = datetime.now()
    target_date = now.date()

    if month and day:
        try:
            target_date = datetime(now.year, int(month), int(day)).date()
        except ValueError:
            target_date = now.date()
    if is_yesterday:
        target_date = (datetime.combine(target_date, datetime.min.time()) - timedelta(days=1)).date()

    try:
        dt = datetime.combine(target_date, datetime.min.time()).replace(hour=int(hour), minute=int(minute))
    except ValueError:
        return None
    return dt


def load_wechat_docx_rows(filepath: str) -> list:
    if not filepath or not os.path.isfile(filepath):
        return []

    lines = list(iter_docx_lines(filepath))
    rows = []
    last_time_text = ""

    for raw in lines:
        text = str(raw or "").strip()
        if not text:
            continue

        last_time_text = parse_docx_time_text(text, last_time_text)
        codes = extract_order_codes_from_text(text)
        if not codes:
            continue

        for code in codes:
            rows.append({
                "chat_time": last_time_text,
                "content": text,
                "order_code": code,
                "base_code": get_base_code(code),
                "source_file": os.path.basename(filepath),
            })

    return rows


def load_wechat_export_rows(filepath: str) -> list:
    """
    读取微信导出：支持 xlsx / docx
    - xlsx: 时间列 B，内容列 E
    - docx: 解析聊天内容
    返回明细行列表（一个内容可能提取多个单号，会拆多行）。
    """
    if not filepath or not os.path.isfile(filepath):
        return []

    suffix = os.path.splitext(filepath)[1].lower()
    if suffix == ".docx":
        return load_wechat_docx_rows(filepath)

    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    ws = wb.active

    rows = []
    for row_idx, row in enumerate(ws.iter_rows(min_row=1, values_only=True), 1):
        time_val = row[1] if len(row) > 1 else ""  # B
        content_val = row[4] if len(row) > 4 else ""  # E

        if is_wechat_noise(content_val):
            continue

        codes = extract_order_codes_from_text(content_val)
        if not codes:
            continue

        time_text = str(time_val or "").strip()
        for code in codes:
            rows.append({
                "chat_time": time_text,
                "content": str(content_val or "").strip(),
                "order_code": code,
                "base_code": get_base_code(code),
                "source_file": os.path.basename(filepath),
            })

    wb.close()
    return rows


def parse_time_for_sort(time_text: str):
    text = str(time_text or "").strip()
    if not text:
        return datetime.min
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M", "%Y/%m/%d %H:%M:%S", "%Y/%m/%d %H:%M"):
        try:
            return datetime.strptime(text, fmt)
        except Exception:
            pass
    return datetime.min


def merge_wechat_rows(existing_rows: list, new_rows: list, dedup_by_order_only: bool = True) -> list:
    """合并微信明细；当前按单号唯一，保留最新消息原文与来源文件。"""
    merged = {}

    def put_row(r):
        code = normalize_order_code(r.get("order_code", ""))
        if not code:
            return
        key = code if dedup_by_order_only else f"{r.get('chat_time','')}|{code}"
        old = merged.get(key)
        if old is None:
            merged[key] = dict(r)
            return
        # 同key时，保留时间更晚的内容作为代表
        if parse_time_for_sort(r.get("chat_time", "")) >= parse_time_for_sort(old.get("chat_time", "")):
            merged[key] = dict(r)

    for row in existing_rows or []:
        put_row(row)
    for row in new_rows or []:
        put_row(row)

    return sorted(
        merged.values(),
        key=lambda x: (x.get("order_code", ""), parse_time_for_sort(x.get("chat_time", ""))),
    )


def aggregate_wechat_rows(rows: list) -> dict:
    """
    聚合为单号索引：
    - first_seen / last_seen / hit_count
    - 同时建立 base_code -> 单号集合
    """
    by_code = {}
    base_index = {}

    for r in rows or []:
        code = normalize_order_code(r.get("order_code", ""))
        if not code:
            continue
        base = get_base_code(code)
        tm = str(r.get("chat_time", "")).strip()

        if code not in by_code:
            by_code[code] = {
                "order_code": code,
                "base_code": base,
                "first_seen": tm,
                "last_seen": tm,
                "hit_count": 1,
            }
        else:
            rec = by_code[code]
            rec["hit_count"] += 1
            if parse_time_for_sort(tm) < parse_time_for_sort(rec["first_seen"]):
                rec["first_seen"] = tm
            if parse_time_for_sort(tm) > parse_time_for_sort(rec["last_seen"]):
                rec["last_seen"] = tm

        if base:
            base_index.setdefault(base, set()).add(code)

    return {"by_code": by_code, "base_index": base_index}


def read_wechat_repo(repo_path: str) -> list:
    """读取汇总库 raw_imports 工作表。"""
    if not repo_path or not os.path.isfile(repo_path):
        return []

    wb = openpyxl.load_workbook(repo_path, read_only=True, data_only=True)
    if "raw_imports" not in wb.sheetnames:
        wb.close()
        return []

    ws = wb["raw_imports"]
    out = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        out.append({
            "chat_time": str(row[0] or "").strip() if len(row) > 0 else "",
            "content": str(row[1] or "").strip() if len(row) > 1 else "",
            "order_code": normalize_order_code(str(row[2] or "").strip() if len(row) > 2 else ""),
            "base_code": str(row[3] or "").strip() if len(row) > 3 else "",
            "source_file": str(row[4] or "").strip() if len(row) > 4 else "",
        })
    wb.close()
    return out


def save_wechat_repo(repo_path: str, raw_rows: list, agg_map: dict):
    """保存微信汇总库：raw_imports + clean_orders。"""
    wb = openpyxl.Workbook()
    ws_raw = wb.active
    ws_raw.title = "raw_imports"

    raw_headers = ["chat_time", "content", "order_code", "base_code", "source_file"]
    for c, h in enumerate(raw_headers, 1):
        ws_raw.cell(row=1, column=c, value=h)

    for i, r in enumerate(raw_rows, 2):
        ws_raw.cell(row=i, column=1, value=r.get("chat_time", ""))
        ws_raw.cell(row=i, column=2, value=r.get("content", ""))
        ws_raw.cell(row=i, column=3, value=r.get("order_code", ""))
        ws_raw.cell(row=i, column=4, value=r.get("base_code", ""))
        ws_raw.cell(row=i, column=5, value=r.get("source_file", ""))

    ws_clean = wb.create_sheet("clean_orders")
    clean_headers = ["order_code", "base_code", "first_seen", "last_seen", "hit_count"]
    for c, h in enumerate(clean_headers, 1):
        ws_clean.cell(row=1, column=c, value=h)

    by_code = agg_map.get("by_code", {})
    sorted_codes = sorted(by_code.keys())
    for i, code in enumerate(sorted_codes, 2):
        rec = by_code[code]
        ws_clean.cell(row=i, column=1, value=rec.get("order_code", ""))
        ws_clean.cell(row=i, column=2, value=rec.get("base_code", ""))
        ws_clean.cell(row=i, column=3, value=rec.get("first_seen", ""))
        ws_clean.cell(row=i, column=4, value=rec.get("last_seen", ""))
        ws_clean.cell(row=i, column=5, value=rec.get("hit_count", 0))

    wb.save(repo_path)


def build_wechat_index(repo_path: str, wechat_export_path: str, mode: str = "append") -> dict:
    """
    构建微信索引：
    - append: 读历史库 + 新导出，合并去重后回写库
    - merge_new: 与 append 行为一致（保留模式开关，便于后续扩展）
    返回 agg_map。
    """
    mode = str(mode or "append").strip().lower()
    existing_rows = read_wechat_repo(repo_path)
    new_rows = load_wechat_export_rows(wechat_export_path)

    if not new_rows:
        return aggregate_wechat_rows(existing_rows)

    merged_rows = merge_wechat_rows(existing_rows, new_rows, dedup_by_order_only=True)
    agg_map = aggregate_wechat_rows(merged_rows)

    # 当前 merge_new 与 append 一致，先满足可切换
    if mode in ("append", "merge_new"):
        save_wechat_repo(repo_path, merged_rows, agg_map)

    return agg_map


def match_wechat_shipping(oid: str, agg_map: dict):
    """
    返回 (has_shipping, match_type, first_seen, last_seen, hit_count)
    match_type: exact / base / none
    """
    code = normalize_order_code(oid)
    base = get_base_code(code)
    by_code = agg_map.get("by_code", {}) if isinstance(agg_map, dict) else {}
    base_index = agg_map.get("base_index", {}) if isinstance(agg_map, dict) else {}

    if code and code in by_code:
        rec = by_code[code]
        return True, "exact", rec.get("first_seen", ""), rec.get("last_seen", ""), rec.get("hit_count", 0)

    if base and base in base_index:
        # 基础号命中时，聚合该基础号下所有子单的时间范围
        matched_codes = sorted(base_index.get(base, []))
        if matched_codes:
            first = ""
            last = ""
            cnt = 0
            for c in matched_codes:
                rec = by_code.get(c, {})
                cnt += int(rec.get("hit_count", 0) or 0)
                r_first = rec.get("first_seen", "")
                r_last = rec.get("last_seen", "")
                if not first or parse_time_for_sort(r_first) < parse_time_for_sort(first):
                    first = r_first
                if not last or parse_time_for_sort(r_last) > parse_time_for_sort(last):
                    last = r_last
            return True, "base", first, last, cnt

    return False, "none", "", "", 0


def final_shipping_status(order_status_text: str, has_shipping: bool, match_type: str) -> str:
    """最终发货判定。"""
    if match_type == "exact" and has_shipping:
        return "已发货"

    is_done = "打包完成" in str(order_status_text or "") or "已完成" in str(order_status_text or "")

    if is_done and has_shipping:
        return "已发货"
    if is_done and not has_shipping:
        return "待核实(已打包未见微信记录)"
    if (not is_done) and has_shipping:
        return "异常(未打包但微信有记录)"
    return "正常未发货"


def filter_exact_orders(orders: list, keyword: str) -> list:
    """
    在接口模糊搜索结果中，保留“基础编号匹配”的订单：
    - keyword=99641 时，99641 / 99641-1 / 99641-2 都应命中
    """
    target = str(keyword or "").strip()
    if not target:
        return []

    exact = []
    for order in orders:
        # 常见可能字段：oId / oid / orderId / no / code
        candidates = [
            order.get("oId"),
            order.get("oid"),
            order.get("orderId"),
            order.get("no"),
            order.get("code"),
        ]

        matched = False
        for v in candidates:
            if v is None:
                continue
            text = str(v).strip()
            if text == target or text.startswith(f"{target}-"):
                matched = True
                break

        if matched:
            exact.append(order)

    return exact


def query_order(keyword: str):
    """
    返回 (found: bool, orders: list)
    found=False → 真的查无此单
    found=True  → 查到订单（可能有多条，如 99641-1 / 99641-2）
    网络失败时自动重试最多3次
    """
    MAX_RETRY = 3
    RETRY_WAIT = 3  # 每次重试前等待秒数

    for attempt in range(1, MAX_RETRY + 1):
        try:
            session = requests.Session()
            session.trust_env = False  # 禁用系统代理，避免代理/证书导致请求失败
            headers = build_headers()

            # 先查总数
            count_resp = session.post(
                API_GET_COUNT,
                json=build_payload(keyword),
                headers=headers,
                timeout=15,
            )

            try:
                count_data = count_resp.json()
            except Exception:
                print(
                    f"    [WARN] 接口返回非JSON（可能被风控/跳转/代理拦截），"
                    f"status={count_resp.status_code}"
                )
                return False, []

            if count_data.get("code") != 200:
                print(
                    f"    [WARN] 接口未返回成功 code=200，"
                    f"code={count_data.get('code')} msg={count_data.get('message', '')}"
                )
                return False, []

            total = int(count_data.get("data", 0))
            if total == 0:
                return False, []  # 确认查无此单

            # 分页拉取订单列表
            page_size = 20
            pages = max(1, -(-total // page_size))
            results = []

            for page in range(1, pages + 1):
                resp = session.post(
                    API_GET_LIST,
                    json=build_payload(keyword, page_index=page, page_size=page_size),
                    headers=headers,
                    timeout=15,
                )
                data = resp.json()
                if data.get("code") != 200:
                    print(f"    ⚠ getList异常: {data.get('message', '')}")
                    break
                orders = data.get("data", [])
                results.extend(orders)
                if len(orders) < page_size:
                    break

            # 关键：基础编号匹配，确保 99641 能匹配 99641-1 / 99641-2
            exact_results = filter_exact_orders(results, keyword)
            if not exact_results:
                return False, []

            # 逐条拉 stepProgress 做状态判定
            for order in exact_results:
                status_text, row_color, process = decide_status_by_step_or_process(session, headers, order)
                order["_status_text"] = status_text
                order["_row_color"] = row_color
                order["_process"] = process

            return True, exact_results

        except requests.exceptions.SSLError as e:
            print(f"    ⚠ SSL/证书错误: {e}")
            return False, []
        except requests.exceptions.ConnectionError:
            print(f"    ⚠ 网络连接失败 (第{attempt}/{MAX_RETRY}次)，{RETRY_WAIT}秒后重试...")
        except requests.exceptions.Timeout:
            print(f"    ⚠ 请求超时 (第{attempt}/{MAX_RETRY}次)，{RETRY_WAIT}秒后重试...")
        except Exception as e:
            print(f"    ⚠ 查询出错 (第{attempt}/{MAX_RETRY}次): {e}，{RETRY_WAIT}秒后重试...")

        if attempt < MAX_RETRY:
            time.sleep(RETRY_WAIT)

    print(f"    ❌ 重试{MAX_RETRY}次均失败，跳过此编号")
    return False, []


def load_input_excel(filepath: str) -> pd.DataFrame:
    df = pd.read_excel(filepath, dtype=str)
    df.columns = df.columns.str.strip()

    col_name, col_id = None, None
    for col in df.columns:
        if "用户" in col:
            col_name = col
        if "编号" in col or "oid" in col.lower():
            col_id = col

    if col_id is None:
        raise RuntimeError(f"找不到订单编号列，当前列名: {list(df.columns)}")

    rename = {col_id: "订单编号"}
    if col_name:
        rename[col_name] = "用户名"
    df = df.rename(columns=rename)

    if "用户名" not in df.columns:
        df["用户名"] = ""

    df = df[df["订单编号"].notna() & (df["订单编号"].str.strip() != "")]
    return df.reset_index(drop=True)


def save_results(all_rows: list, filepath: str, include_wechat: bool = False):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "查询结果"

    thin = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    header_font = Font(name="Arial", bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", start_color="4472C4")
    normal_font = Font(name="Arial", size=10)

    headers = [
        "用户名",
        "原始输入",
        "查询编号",
        "订单编号(oId)",
        "订单内部ID(id)",
        "产品名称",
        "颜色/规格",
        "总数量",
        "订单状态",
        "进度百分比",
    ]
    if include_wechat:
        headers += ["微信匹配", "微信首现", "微信末现", "微信次数", "发货判定"]
    headers += ["下单时间", "查询时间"]

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border
    ws.row_dimensions[1].height = 22

    for row_idx, row in enumerate(all_rows, 2):
        row_fill = PatternFill("solid", start_color=row.get("_color", COLOR_NOT_FOUND))
        values = [
            row.get("用户名", ""),
            row.get("原始输入", ""),
            row.get("查询编号", ""),
            row.get("oId", ""),
            row.get("id", ""),
            row.get("name", ""),
            row.get("color_size", ""),
            row.get("totalQuantity", ""),
            row.get("状态文字", ""),
            row.get("进度百分比", ""),
        ]
        if include_wechat:
            values += [
                row.get("微信匹配", ""),
                row.get("微信首现", ""),
                row.get("微信末现", ""),
                row.get("微信次数", ""),
                row.get("发货判定", ""),
            ]
        values += [
            row.get("orderTime", ""),
            row.get("查询时间", ""),
        ]
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.font = normal_font
            cell.alignment = center
            cell.border = border
            cell.fill = row_fill
        ws.row_dimensions[row_idx].height = 18

    if include_wechat:
        col_widths = [12, 22, 14, 14, 16, 22, 16, 10, 14, 12, 10, 17, 17, 10, 22, 18, 18]
    else:
        col_widths = [12, 22, 14, 14, 16, 22, 16, 10, 14, 12, 18, 18]
    for col, w in enumerate(col_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = w
    ws.freeze_panes = "A2"

    # 图例说明页
    ws2 = wb.create_sheet("颜色说明")
    legend = [
        (COLOR_DONE,      "✅ 绿色", "已完成（进度≥75% 或 stepProgress命中打包完成）"),
        (COLOR_IN_PROG,   "🟡 黄色", "生产进行中（0%<进度<75%）"),
        (COLOR_PENDING,   "🔵 蓝色", "待开始（订单存在，尚未开始生产）"),
        (COLOR_NOT_FOUND, "🔴 红色", "查无此单（编号不存在）"),
    ]
    for i, (color, label, desc) in enumerate(legend, 1):
        fill = PatternFill("solid", start_color=color)
        for col_idx, val in enumerate([label, desc], 1):
            c = ws2.cell(row=i, column=col_idx, value=val)
            c.fill = fill
            c.font = Font(name="Arial", size=10)
            c.alignment = Alignment(horizontal="left", vertical="center")
        ws2.row_dimensions[i].height = 22
    ws2.column_dimensions["A"].width = 12
    ws2.column_dimensions["B"].width = 38

    wb.save(filepath)
    print(f"✅ 结果已保存: {filepath}")


def parse_cli_args():
    parser = argparse.ArgumentParser(description="棉花果订单批量查询")
    parser.add_argument("--input", default="", help="输入Excel路径")
    parser.add_argument("--output", default="", help="输出Excel路径")
    parser.add_argument("--cookie", default="", help="Cookie")
    parser.add_argument("--start-date", default="", help="开始日期，如 2024-01-01")
    parser.add_argument("--end-date", default="", help="结束日期，如 2026-12-31")
    parser.add_argument("--delay", type=float, default=None, help="请求间隔秒数")
    parser.add_argument("--wechat", default="", help="微信导出xlsx路径")
    parser.add_argument("--wechat-log", default="", help="微信汇总库路径")
    parser.add_argument("--wechat-mode", default="", choices=["append", "merge_new"], help="微信导入模式")
    parser.add_argument("--enable-wechat-compare", action="store_true", help="启用微信比对（默认关闭）")
    parser.add_argument("--no-interactive", action="store_true", help="非交互模式，使用参数/配置直接运行")
    return parser.parse_args()


def apply_cli_args(args):
    if args.cookie:
        CONFIG["cookie"] = args.cookie
    if args.start_date:
        CONFIG["start_date"] = args.start_date
    if args.end_date:
        CONFIG["end_date"] = args.end_date
    if args.delay is not None:
        CONFIG["request_delay"] = max(0.0, float(args.delay))
    if args.wechat:
        CONFIG["wechat_export_excel"] = args.wechat
    if args.wechat_log:
        CONFIG["wechat_repo_excel"] = args.wechat_log
    if args.wechat_mode:
        CONFIG["wechat_mode"] = args.wechat_mode


def main():
    init_console_encoding()
    print("=" * 55)
    print("  棉花果订单批量查询爬虫  v2.2")
    print("=" * 55)
    print("[STEP 0/5] 初始化")

    args = parse_cli_args()
    apply_cli_args(args)

    if not CONFIG["cookie"] or len(CONFIG["cookie"]) < 10:
        print("\n❌ 请先在 CONFIG 中填入你的 Cookie 或通过 --cookie 传入！")
        return

    if args.no_interactive:
        input_excel = clean_cli_path(args.input) if args.input else clean_cli_path(CONFIG["input_excel"])
        output_excel = clean_cli_path(args.output) if args.output else clean_cli_path(CONFIG["output_excel"])
        wechat_excel = clean_cli_path(args.wechat) if args.wechat else clean_cli_path(CONFIG.get("wechat_export_excel", ""))
        wechat_repo = clean_cli_path(args.wechat_log) if args.wechat_log else clean_cli_path(CONFIG.get("wechat_repo_excel", "wechat_shipment_log.xlsx"))
        wechat_mode = args.wechat_mode if args.wechat_mode else CONFIG.get("wechat_mode", "append")
    else:
        input_excel, output_excel, wechat_excel, wechat_repo, wechat_mode = prompt_runtime_paths()

    wechat_agg_map = {"by_code": {}, "base_index": {}}
    enable_wechat_compare = bool(args.enable_wechat_compare)

    if enable_wechat_compare:
        if wechat_excel:
            print(f"\n[STEP 1/5] 导入微信导出表: {wechat_excel}")
            print(f"[INFO] 微信汇总库: {wechat_repo}  模式: {wechat_mode}")
            try:
                wechat_agg_map = build_wechat_index(wechat_repo, wechat_excel, mode=wechat_mode)
                print(f"[OK] 微信汇总库已更新: {wechat_repo}  模式: {wechat_mode}")
                print(f"[INFO] 已收录唯一单号: {len(wechat_agg_map.get('by_code', {}))}")
            except Exception as e:
                print(f"[ERR] 微信导出处理失败: {e}")
                return
        else:
            print("\n[ERR] 已启用微信比对，但未提供微信导出表。请检查 --wechat 参数或GUI文件选择。")
            return
    else:
        print("\n[INFO] 未启用微信比对（默认关闭），仅输出订单进度。")

    print(f"\n[STEP 2/5] 读取Excel: {input_excel}")
    try:
        df = load_input_excel(input_excel)
    except RuntimeError as e:
        print(f"❌ {e}")
        return

    total_records = len(df)
    print(f"[INFO] 共读取到 {total_records} 条查询记录\n")

    all_rows = []
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cnt_done = cnt_pending = cnt_inprog = cnt_notfound = 0

    print("[STEP 3/5] 查询订单并判定状态")

    for idx, row in df.iterrows():
        username = str(row.get("用户名", "")).strip()
        if username in ("nan", "None", ""):
            username = ""
        raw_keyword = str(row["订单编号"]).strip()
        keyword = normalize_query_id(raw_keyword)

        print(f"[{idx+1}/{total_records}] 原始: {raw_keyword} -> 查询编号: {keyword}  用户: {username or '(无)'}")

        found, orders = query_order(keyword)

        if not found:
            cnt_notfound += 1
            print(f"    → ❌ 查无此单")
            all_rows.append({
                "用户名": username,
                "原始输入": raw_keyword,
                "查询编号": keyword,
                "oId": "",
                "id": "",
                "name": "查无此单",
                "color_size": "", "totalQuantity": "",
                "状态文字": "查无此单", "进度百分比": "",
                "微信匹配": "none",
                "微信首现": "",
                "微信末现": "",
                "微信次数": 0,
                "发货判定": "查无此单",
                "orderTime": "", "查询时间": now_str,
                "_color": COLOR_NOT_FOUND,
            })
        else:
            for order in orders:
                process = parse_progress_value(order.get("_process", order.get("process")))
                status_text = order.get("_status_text", "") or process_to_status(process)[0]
                row_color = order.get("_row_color", "") or process_to_status(process)[1]

                # 统计：若命中“打包完成”则优先计为已完成
                if "打包完成" in status_text:
                    cnt_done += 1
                elif process >= 0.75:
                    cnt_done += 1
                elif process == 0.0:
                    cnt_pending += 1
                else:
                    cnt_inprog += 1

                # 颜色规格汇总
                quantities = order.get("quantities", {})
                color_size_parts = [
                    f"{c}/{s}×{q}"
                    for c, sizes in quantities.items()
                    for s, q in sizes.items()
                ]
                color_size = "  ".join(color_size_parts)

                # 时间处理
                order_ts = order.get("orderTime")
                order_time = ""
                if order_ts:
                    try:
                        order_time = datetime.fromtimestamp(order_ts / 1000).strftime("%Y-%m-%d %H:%M")
                    except:
                        pass

                oid = str(order.get("oId", "")).strip()
                if enable_wechat_compare:
                    has_ship, match_type, first_seen, last_seen, hit_count = match_wechat_shipping(oid, wechat_agg_map)
                    ship_status = final_shipping_status(status_text, has_ship, match_type)
                else:
                    match_type, first_seen, last_seen, hit_count = "N/A", "", "", ""
                    ship_status = "未启用微信比对"

                all_rows.append({
                    "用户名": username,
                    "原始输入": raw_keyword,
                    "查询编号": keyword,
                    "oId": oid,
                    "id": order.get("id", ""),
                    "name": order.get("name", ""),
                    "color_size": color_size,
                    "totalQuantity": order.get("totalQuantity", ""),
                    "状态文字": status_text,
                    "进度百分比": f"{int(process*100)}%",
                    "微信匹配": match_type,
                    "微信首现": first_seen,
                    "微信末现": last_seen,
                    "微信次数": hit_count,
                    "发货判定": ship_status,
                    "orderTime": order_time,
                    "查询时间": now_str,
                    "_color": row_color,
                })
                print(
                    f"    → {order.get('name','')} | {status_text} | 微信:{match_type} | 发货判定:{ship_status} | 数量:{order.get('totalQuantity','')}"
                )

        time.sleep(CONFIG["request_delay"])

    print(f"\n[STEP 4/5] 保存结果...")
    save_results(all_rows, output_excel, include_wechat=enable_wechat_compare)

    print(f"\n[STEP 5/5] 查询完成")
    print(f"{'='*55}")
    print(f"  查询完成！共处理 {total_records} 个编号，{len(all_rows)} 条记录")
    print(f"  [OK] 已完成:   {cnt_done} 条")
    print(f"  [INFO] 待开始:   {cnt_pending} 条  ← 订单存在但尚未生产")
    print(f"  [INFO] 进行中:   {cnt_inprog} 条")
    print(f"  [ERR] 查无此单: {cnt_notfound} 条")
    print(f"{'='*55}")


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"\n❌ 运行异常: {e}")

    if "--no-interactive" not in sys.argv:
        try:
            input("\n按回车退出...")
        except Exception:
            pass
