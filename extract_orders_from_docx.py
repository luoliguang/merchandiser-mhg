"""
从 DOCX 文档中提取订单号，输出为与扫描脚本一致的 Excel 格式。

依赖：pip install python-docx openpyxl
"""

import os
import re
import time
import unicodedata
import zipfile
from datetime import datetime
from typing import Iterable, List

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from scan_orders_main import parse_order_base_id, save_excel


ORDER_ID_REGEX = re.compile(r"(?<!\d)[1-9]\d{4,7}(?:-\d{1,3}){0,3}(?!\d)")


def iter_docx_lines(doc_path: str) -> Iterable[str]:
    if not zipfile.is_zipfile(doc_path):
        raise ValueError(
            f"文件不是有效的 DOCX（zip 容器）: {doc_path}\n"
            "请确认文件后缀为 .docx，且不是 .doc 或已损坏文件。"
        )
    try:
        doc = Document(doc_path)
    except PackageNotFoundError as exc:
        raise FileNotFoundError(f"无法打开 DOCX 文件: {doc_path}") from exc

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            yield text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if text:
                    yield text


def normalize_line(text: str) -> str:
    cleaned = text.replace("/", " ").replace("，", " ").replace(",", " ")
    cleaned = cleaned.replace("；", " ").replace(";", " ")
    return cleaned


def extract_order_ids(lines: Iterable[str]) -> List[str]:
    seen = set()
    ordered = []

    for raw in lines:
        line = normalize_line(raw)
        for match in ORDER_ID_REGEX.findall(line):
            if match in seen:
                continue
            seen.add(match)
            ordered.append(match)

    return ordered


def build_order_rows(order_ids: List[str]):
    rows = []
    for order_id in order_ids:
        base_id = parse_order_base_id(order_id)
        rows.append(
            {
                "用户名": "",
                "订单编号": order_id,
                "基础编号": base_id if base_id is not None else "",
                "状态": "已找到",
            }
        )
    return rows


def normalize_input_path(text: str) -> str:
    cleaned = (text or "").strip().strip("\"").strip("'")
    cleaned = unicodedata.normalize("NFKC", cleaned)
    cleaned = cleaned.replace("\\ ", "\\").replace("/ ", "/")
    cleaned = re.sub(r"([\\/])[\s\u3000\u00a0]+", r"\1", cleaned)
    cleaned = re.sub(r"[\u200b\ufeff]", "", cleaned)
    cleaned = cleaned.strip()

    if cleaned and not os.path.isfile(cleaned):
        dir_name, base_name = os.path.split(cleaned)
        if dir_name:
            base_stripped = re.sub(r"\s+", "", base_name)
            candidate = os.path.join(dir_name, base_stripped)
            if os.path.isfile(candidate):
                return candidate

    return cleaned


def resolve_path_by_name(folder: str, target_name: str) -> str:
    if not folder or not os.path.isdir(folder):
        return ""
    normalized_target = re.sub(r"\s+", "", unicodedata.normalize("NFKC", target_name))
    for name in os.listdir(folder):
        if re.sub(r"\s+", "", unicodedata.normalize("NFKC", name)) == normalized_target:
            candidate = os.path.join(folder, name)
            if os.path.isfile(candidate):
                return candidate
    return ""


def resolve_docx_path(raw_input: str) -> str:
    cleaned = normalize_input_path(raw_input)
    if cleaned and os.path.isfile(cleaned):
        return cleaned

    if cleaned:
        folder = os.path.dirname(cleaned)
        name = os.path.basename(cleaned)
        if folder and name:
            by_name = resolve_path_by_name(folder, name)
            if by_name:
                return by_name

    data_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data")
    if os.path.isdir(data_dir):
        docx_files = [
            os.path.join(data_dir, f)
            for f in os.listdir(data_dir)
            if f.lower().endswith(".docx")
        ]
        if docx_files:
            return max(docx_files, key=os.path.getmtime)

    return ""


def main():
    print("= DOCX 订单号提取工具 =")
    raw = input("DOCX 路径（支持拖拽，回车自动识别 data 目录最新文件）：\n> ")
    doc_path = resolve_docx_path(raw)

    if not doc_path or not os.path.isfile(doc_path):
        print("[错误] DOCX 文件不存在，且未在 data 目录找到可用文件。")
        return

    default_output = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "data",
        f"orders_result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
    )
    output_path = input(f"输出 Excel（回车使用默认）\n> {default_output}\n> ").strip()
    if not output_path:
        output_path = default_output

    start = time.time()
    lines = list(iter_docx_lines(doc_path))
    order_ids = extract_order_ids(lines)
    rows = build_order_rows(order_ids)

    save_excel(rows, output_path)

    elapsed = time.time() - start
    print(f"\n完成：共提取 {len(order_ids)} 个订单号")
    print(f"输出文件：{output_path}")
    print(f"耗时：{elapsed:.2f}s")


if __name__ == "__main__":
    main()
